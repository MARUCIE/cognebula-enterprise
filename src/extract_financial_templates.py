#!/usr/bin/env python3
"""
Extract structured content from the 12-chapter enterprise financial management system.

Walks through all chapters under:
  doc-tax/财税脑图3.0/3.财税实战工具大全/4.行业财务制度范本/企业实用管理制度表格范本模版【全套】/

Outputs:
  - data/extracted/financial_templates/compliance_rules.json   (from 制度/ files)
  - data/extracted/financial_templates/form_templates.json      (from 表格/ files)
  - data/extracted/financial_templates/manifest.json            (summary)
"""

import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import docx


# -- Constants --

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SOURCE_BASE = PROJECT_ROOT / (
    "doc-tax/财税脑图3.0/3.财税实战工具大全/4.行业财务制度范本/"
    "企业实用管理制度表格范本模版【全套】"
)
OUTPUT_DIR = PROJECT_ROOT / "data/extracted/financial_templates"

CHAPTER_MAP = {
    "第1章": {"number": 1, "name": "总目录", "category": "overview"},
    "第2章": {"number": 2, "name": "货币资金管理", "category": "cash_management"},
    "第3章": {"number": 3, "name": "往来账管理", "category": "receivables_payables"},
    "第4章": {"number": 4, "name": "资产管理", "category": "asset_management"},
    "第5章": {"number": 5, "name": "成本费用管理", "category": "cost_expense"},
    "第6章": {"number": 6, "name": "收入利润核算管理", "category": "revenue_profit"},
    "第7章": {"number": 7, "name": "税务管理", "category": "tax_management"},
    "第8章": {"number": 8, "name": "会计资料管理", "category": "accounting_records"},
    "第9章": {"number": 9, "name": "投融资管理", "category": "investment_financing"},
    "第10章": {"number": 10, "name": "预算管理", "category": "budget_management"},
    "第11章": {"number": 11, "name": "财务部日常管理", "category": "finance_dept_ops"},
    "第12章": {"number": 12, "name": "审计管理", "category": "audit_management"},
}


# -- Extraction helpers --

def extract_paragraphs(doc_path: Path) -> list[str]:
    """Extract all non-empty paragraph texts from a docx file."""
    try:
        doc = docx.Document(str(doc_path))
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        print(f"  WARN: Failed to read paragraphs from {doc_path.name}: {e}")
        return []


def extract_tables(doc_path: Path) -> list[list[list[str]]]:
    """Extract all tables as list of rows (each row = list of cell texts)."""
    try:
        doc = docx.Document(str(doc_path))
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(rows)
        return tables
    except Exception as e:
        print(f"  WARN: Failed to read tables from {doc_path.name}: {e}")
        return []


def parse_chapter_key(dirname: str) -> str | None:
    """Extract chapter key like '第2章' from directory name like '第2章 货币资金管理'."""
    m = re.match(r"(第\d+章)", dirname)
    return m.group(1) if m else None


def extract_key_requirements(paragraphs: list[str]) -> list[str]:
    """Pull out numbered articles (第X条) as key requirements."""
    requirements = []
    for p in paragraphs:
        # Match lines starting with 第X条 pattern
        if re.match(r"第[一二三四五六七八九十百千\d]+条", p):
            requirements.append(p)
    return requirements


def extract_form_fields(tables: list[list[list[str]]]) -> list[str]:
    """Extract unique field names from table headers (first row or first column)."""
    fields = []
    seen = set()
    for table in tables:
        if not table:
            continue
        # Collect from first row (header)
        for cell in table[0]:
            cleaned = cell.strip()
            if cleaned and cleaned not in seen and len(cleaned) < 50:
                fields.append(cleaned)
                seen.add(cleaned)
        # Collect from first column (row labels)
        for row in table[1:]:
            if row:
                cleaned = row[0].strip()
                if cleaned and cleaned not in seen and len(cleaned) < 50:
                    fields.append(cleaned)
                    seen.add(cleaned)
    return fields


def infer_purpose(filename: str, paragraphs: list[str]) -> str:
    """Infer form purpose from filename and first meaningful paragraph."""
    name = filename.replace(".docx", "").replace(".doc", "")
    # Use first non-title paragraph if available
    for p in paragraphs:
        if len(p) > 20 and not re.match(r"^[\s\d.、]+$", p):
            return p[:200]
    return name


# -- Main extraction --

def process_rule_file(filepath: Path, chapter_info: dict) -> dict:
    """Process a 制度 file into a ComplianceRule node."""
    name = filepath.stem
    paragraphs = extract_paragraphs(filepath)
    key_reqs = extract_key_requirements(paragraphs)
    full_text = "\n".join(paragraphs)

    return {
        "node_type": "ComplianceRule",
        "name": name,
        "chapter_number": chapter_info["number"],
        "chapter_name": chapter_info["name"],
        "category": chapter_info["category"],
        "source_file": filepath.name,
        "source_path": str(filepath.relative_to(PROJECT_ROOT)),
        "article_count": len(key_reqs),
        "key_requirements": key_reqs[:20],  # Cap at 20 to keep manageable
        "full_text": full_text,
        "char_count": len(full_text),
    }


def process_form_file(filepath: Path, chapter_info: dict) -> dict:
    """Process a 表格 file into a FormTemplate node."""
    name = filepath.stem
    paragraphs = extract_paragraphs(filepath)
    tables = extract_tables(filepath)
    fields = extract_form_fields(tables)
    purpose = infer_purpose(filepath.name, paragraphs)

    # Serialize tables to a compact representation
    table_data = []
    for table in tables:
        table_data.append({"rows": len(table), "cols": len(table[0]) if table else 0, "data": table})

    return {
        "node_type": "FormTemplate",
        "name": name,
        "chapter_number": chapter_info["number"],
        "chapter_name": chapter_info["name"],
        "category": chapter_info["category"],
        "source_file": filepath.name,
        "source_path": str(filepath.relative_to(PROJECT_ROOT)),
        "fields": fields,
        "field_count": len(fields),
        "table_count": len(tables),
        "tables": table_data,
        "purpose": purpose,
        "paragraphs": paragraphs,
    }


def run_extraction():
    """Walk all 12 chapters and extract structured content."""
    if not SOURCE_BASE.exists():
        print(f"ERROR: Source directory not found: {SOURCE_BASE}")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    compliance_rules = []
    form_templates = []
    stats = {
        "chapters_processed": 0,
        "rule_files": 0,
        "form_files": 0,
        "rule_files_failed": 0,
        "form_files_failed": 0,
        "total_articles": 0,
        "total_fields": 0,
        "per_chapter": {},
    }

    # Sort chapters by number
    chapter_dirs = sorted(SOURCE_BASE.iterdir(), key=lambda d: d.name)

    for chapter_dir in chapter_dirs:
        if not chapter_dir.is_dir():
            continue

        chapter_key = parse_chapter_key(chapter_dir.name)
        if not chapter_key or chapter_key not in CHAPTER_MAP:
            print(f"  NOTE: Skipping unrecognized directory: {chapter_dir.name}")
            continue

        chapter_info = CHAPTER_MAP[chapter_key]
        ch_num = chapter_info["number"]
        print(f"\n=== Chapter {ch_num}: {chapter_info['name']} ===")

        ch_stats = {"rules": 0, "forms": 0, "articles": 0, "fields": 0}

        # Process 制度/ (rules/procedures)
        rules_dir = chapter_dir / "制度"
        if rules_dir.exists():
            for f in sorted(rules_dir.glob("*.docx")):
                print(f"  [Rule] {f.name}")
                try:
                    rule = process_rule_file(f, chapter_info)
                    compliance_rules.append(rule)
                    stats["rule_files"] += 1
                    ch_stats["rules"] += 1
                    ch_stats["articles"] += rule["article_count"]
                    stats["total_articles"] += rule["article_count"]
                except Exception as e:
                    print(f"  ERROR: {f.name}: {e}")
                    stats["rule_files_failed"] += 1

        # Process 表格/ (forms)
        forms_dir = chapter_dir / "表格"
        if forms_dir.exists():
            for f in sorted(forms_dir.glob("*.docx")):
                print(f"  [Form] {f.name}")
                try:
                    form = process_form_file(f, chapter_info)
                    form_templates.append(form)
                    stats["form_files"] += 1
                    ch_stats["forms"] += 1
                    ch_stats["fields"] += form["field_count"]
                    stats["total_fields"] += form["field_count"]
                except Exception as e:
                    print(f"  ERROR: {f.name}: {e}")
                    stats["form_files_failed"] += 1

        stats["per_chapter"][f"ch{ch_num}_{chapter_info['category']}"] = ch_stats
        stats["chapters_processed"] += 1

    # Write outputs
    rules_path = OUTPUT_DIR / "compliance_rules.json"
    forms_path = OUTPUT_DIR / "form_templates.json"
    manifest_path = OUTPUT_DIR / "manifest.json"

    with open(rules_path, "w", encoding="utf-8") as f:
        json.dump(compliance_rules, f, ensure_ascii=False, indent=2)
    print(f"\nOK: Wrote {len(compliance_rules)} ComplianceRule nodes -> {rules_path.relative_to(PROJECT_ROOT)}")

    with open(forms_path, "w", encoding="utf-8") as f:
        json.dump(form_templates, f, ensure_ascii=False, indent=2)
    print(f"OK: Wrote {len(form_templates)} FormTemplate nodes -> {forms_path.relative_to(PROJECT_ROOT)}")

    manifest = {
        "extraction_date": datetime.now().isoformat(),
        "source": str(SOURCE_BASE.relative_to(PROJECT_ROOT)),
        "output_dir": str(OUTPUT_DIR.relative_to(PROJECT_ROOT)),
        "files": {
            "compliance_rules": str(rules_path.relative_to(PROJECT_ROOT)),
            "form_templates": str(forms_path.relative_to(PROJECT_ROOT)),
        },
        "stats": stats,
        "node_types": {
            "ComplianceRule": {
                "count": len(compliance_rules),
                "schema": ["name", "chapter_number", "chapter_name", "category",
                           "source_file", "article_count", "key_requirements", "full_text"],
            },
            "FormTemplate": {
                "count": len(form_templates),
                "schema": ["name", "chapter_number", "chapter_name", "category",
                           "source_file", "fields", "field_count", "tables", "purpose"],
            },
        },
    }

    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    print(f"OK: Wrote manifest -> {manifest_path.relative_to(PROJECT_ROOT)}")

    # Summary
    print(f"\n{'='*60}")
    print(f"EXTRACTION COMPLETE")
    print(f"  Chapters processed: {stats['chapters_processed']}")
    print(f"  ComplianceRule nodes: {stats['rule_files']} (failed: {stats['rule_files_failed']})")
    print(f"  FormTemplate nodes: {stats['form_files']} (failed: {stats['form_files_failed']})")
    print(f"  Total articles extracted: {stats['total_articles']}")
    print(f"  Total form fields extracted: {stats['total_fields']}")
    print(f"{'='*60}")


if __name__ == "__main__":
    run_extraction()
