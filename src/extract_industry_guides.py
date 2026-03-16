#!/usr/bin/env python3
"""Extract industry-specific accounting/tax guides from doc-tax.

Processes doc/docx/pdf files from 行业账务风险筹划 directory.
Maps each file to industry + content type for graph enrichment:
- OP_BusinessScenario nodes (industry-specific accounting flows)
- RiskIndicator nodes (industry tax risk points)
- OP_StandardCase nodes (typical journal entries per industry)

Usage:
    python3 src/extract_industry_guides.py
"""

import json
import subprocess
import sys
from pathlib import Path

BASE_DIR = Path("doc-tax/财税脑图3.0/3.财税实战工具大全/5.行业账务风险筹划")
OUTPUT_DIR = Path("data/extracted/industry_guides")

# Industry classification from file names
INDUSTRY_MAP = {
    "建筑": "construction",
    "房地产": "real_estate",
    "房开": "real_estate",
    "汽车": "auto_dealer",
    "4S店": "auto_dealer",
    "医美": "medical_aesthetics",
    "医疗美容": "medical_aesthetics",
    "物业": "property_mgmt",
    "网络直播": "live_streaming",
    "直播": "live_streaming",
    "软件": "software",
    "高新": "high_tech",
    "采矿": "mining",
    "煤炭": "coal_mining",
    "再生资源": "recycling",
    "幼儿园": "kindergarten",
    "律师": "law_firm",
    "人力资源": "hr_services",
    "劳务派遣": "labor_dispatch",
    "合伙企业": "partnership",
    "非营利": "nonprofit",
    "制造": "manufacturing",
    "出口退税": "export_tax_refund",
    "混凝土": "concrete",
    "黄金": "gold_retail",
    "跨境电商": "cross_border_ecommerce",
}

# Content type classification
CONTENT_TYPE_MAP = {
    "会计分录": "journal_entries",
    "账务处理": "accounting_treatment",
    "税收检查": "tax_audit_guide",
    "税务风险": "tax_risk_analysis",
    "财税合规": "compliance_guide",
    "税筹": "tax_planning",
    "会计制度": "accounting_standards",
    "税收优惠": "tax_incentives",
    "涉税风险": "tax_risk_analysis",
    "核算手册": "accounting_manual",
    "做账": "bookkeeping_flow",
}


def classify_file(filename: str) -> dict:
    """Classify a file by industry and content type from its name."""
    industry = "general"
    content_type = "general_guide"

    for cn_key, en_val in INDUSTRY_MAP.items():
        if cn_key in filename:
            industry = en_val
            break

    for cn_key, en_val in CONTENT_TYPE_MAP.items():
        if cn_key in filename:
            content_type = en_val
            break

    return {"industry": industry, "content_type": content_type}


def extract_doc(filepath: Path) -> str:
    """Extract text from .doc file using macOS textutil."""
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(filepath)],
            capture_output=True, text=True, timeout=30
        )
        return result.stdout.strip()
    except Exception as e:
        return f"[extraction_error: {e}]"


def extract_docx(filepath: Path) -> str:
    """Extract text from .docx file using python-docx."""
    try:
        import docx
        doc = docx.Document(str(filepath))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                prefix = ""
                if para.style and para.style.name:
                    if "Heading" in para.style.name:
                        level = para.style.name.replace("Heading ", "").strip()
                        prefix = f"[H{level}] "
                paragraphs.append(f"{prefix}{para.text.strip()}")
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[extraction_error: {e}]"


def extract_pdf(filepath: Path) -> str:
    """Extract text from PDF using PyMuPDF if available, else pdfplumber."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(filepath))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts).strip()
    except ImportError:
        pass
    try:
        import pdfplumber
        with pdfplumber.open(str(filepath)) as pdf:
            text_parts = []
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts).strip()
    except ImportError:
        return "[no_pdf_library_available]"
    except Exception as e:
        return f"[extraction_error: {e}]"


def extract_key_sections(text: str) -> dict:
    """Extract key sections from Chinese accounting/tax text."""
    sections = {
        "journal_entries": [],    # 会计分录
        "tax_rates": [],          # 税率
        "risk_points": [],        # 风险点
        "compliance_items": [],   # 合规事项
        "filing_steps": [],       # 申报步骤
    }

    lines = text.split("\n")
    current_section = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect journal entry patterns: 借/贷
        if ("借：" in line or "借:" in line or "贷：" in line or "贷:" in line
                or line.startswith("借") or line.startswith("贷")):
            sections["journal_entries"].append(line)
            current_section = "journal_entries"
        # Detect tax rate patterns
        elif any(kw in line for kw in ["税率", "%", "‰", "征收率"]):
            sections["tax_rates"].append(line)
        # Detect risk indicators
        elif any(kw in line for kw in ["风险", "预警", "异常", "稽查", "违规"]):
            sections["risk_points"].append(line)
        # Detect compliance items
        elif any(kw in line for kw in ["合规", "备案", "留存", "资料", "证明"]):
            sections["compliance_items"].append(line)
        # Detect filing steps
        elif any(kw in line for kw in ["申报", "填报", "纳税", "扣缴"]):
            sections["filing_steps"].append(line)
        elif current_section == "journal_entries" and line:
            # Continue capturing journal entry context
            if any(kw in line for kw in ["借", "贷", "应交", "银行", "库存", "主营"]):
                sections["journal_entries"].append(line)
            else:
                current_section = None

    return sections


def process_all():
    """Process all files in the industry guides directory."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    results = []
    skipped = []

    for filepath in sorted(BASE_DIR.iterdir()):
        if filepath.is_dir():
            continue

        suffix = filepath.suffix.lower()
        filename = filepath.name

        classification = classify_file(filename)
        print(f"Processing: {filename} -> {classification['industry']}/{classification['content_type']}")

        # Extract text based on file type
        if suffix == ".doc":
            text = extract_doc(filepath)
        elif suffix == ".docx":
            text = extract_docx(filepath)
        elif suffix == ".pdf":
            text = extract_pdf(filepath)
        elif suffix == ".pptx":
            # Skip pptx for now (need python-pptx)
            skipped.append(filename)
            continue
        else:
            skipped.append(filename)
            continue

        if not text or text.startswith("["):
            skipped.append(filename)
            continue

        # Extract key sections
        sections = extract_key_sections(text)

        # Build result
        entry = {
            "source_file": filename,
            "industry": classification["industry"],
            "content_type": classification["content_type"],
            "title": filepath.stem,
            "text_length": len(text),
            "sections": {
                k: v[:50] for k, v in sections.items()  # Limit for JSON size
            },
            "section_counts": {k: len(v) for k, v in sections.items()},
            "first_500_chars": text[:500],
        }

        results.append(entry)

        # Save individual file extraction
        individual_out = OUTPUT_DIR / f"{classification['industry']}_{filepath.stem[:60]}.json"
        with open(individual_out, "w", encoding="utf-8") as f:
            json.dump({
                **entry,
                "full_text": text[:10000],  # Cap at 10K chars for JSON
                "sections": sections,
            }, f, ensure_ascii=False, indent=2)

    # Save manifest
    manifest = {
        "total_processed": len(results),
        "total_skipped": len(skipped),
        "skipped_files": skipped,
        "industries": list(set(r["industry"] for r in results)),
        "content_types": list(set(r["content_type"] for r in results)),
        "files": results,
    }

    manifest_path = OUTPUT_DIR / "_manifest.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    print(f"\n=== Summary ===")
    print(f"Processed: {len(results)} files")
    print(f"Skipped: {len(skipped)} files ({skipped})")
    print(f"Industries: {manifest['industries']}")
    print(f"Content types: {manifest['content_types']}")

    # Print section stats
    total_entries = sum(r["section_counts"]["journal_entries"] for r in results)
    total_risks = sum(r["section_counts"]["risk_points"] for r in results)
    total_compliance = sum(r["section_counts"]["compliance_items"] for r in results)
    print(f"Journal entries found: {total_entries}")
    print(f"Risk points found: {total_risks}")
    print(f"Compliance items found: {total_compliance}")

    return manifest


if __name__ == "__main__":
    manifest = process_all()
