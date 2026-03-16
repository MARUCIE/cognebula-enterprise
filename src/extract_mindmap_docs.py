#!/usr/bin/env python3
"""Extract structured content from 财税脑图3.0 mindmap documents for knowledge graph ingestion.

Walks the 财税脑图3.0/1.财税脑图分类合集 directory, extracts heading-based structure
from docx files and plain text from .doc files, then maps content to L2 (操作层)
and L3 (合规层) graph node types:

  - OP_BusinessScenario: tax audit scenarios, organization rules, company rules
  - OP_StandardCase: accounting treatment guides, invoice rules, cost/expense rules
  - RiskIndicator: risk warnings, audit triggers
  - TaxOptimization: tax optimization strategies (custom, for later mapping)

Output: structured JSON files in data/extracted/mindmap/

Usage:
    python src/extract_mindmap_docs.py
    python src/extract_mindmap_docs.py --base-dir /path/to/cognebula-enterprise
"""

import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional


# ── Constants ────────────────────────────────────────────────────────────

MINDMAP_SUBDIR = "doc-tax/财税脑图3.0/1.财税脑图分类合集"

# Directory name -> graph node type mapping
CATEGORY_NODE_MAP = {
    "税务稽查": "OP_BusinessScenario",
    "组织相关": "OP_BusinessScenario",
    "公司相关": "OP_BusinessScenario",
    "账务处理": "OP_StandardCase",
    "发票相关": "OP_StandardCase",
    "成本费用": "OP_StandardCase",
    "财税优化": "TaxOptimization",
}

# Top-level standalone files (directly under 分类合集/)
STANDALONE_NODE_MAP = {
    "税费相关": "OP_BusinessScenario",
    "行业相关": "OP_BusinessScenario",
}

# Sub-category keywords -> finer node type override
RISK_KEYWORDS = [
    "风险", "稽查", "预警", "异常", "违规", "处罚", "罚款",
    "易错", "常见问题", "注意事项",
]


# ── Heading-level detection ──────────────────────────────────────────────

def get_heading_level(style_name: str) -> int:
    """Return heading level (1-9) from style name, or 0 for non-heading."""
    if not style_name:
        return 0
    m = re.match(r"[Hh]eading\s*(\d)", style_name)
    if m:
        return int(m.group(1))
    # Chinese custom styles like "样式1" used in some files
    m2 = re.match(r"样式(\d)", style_name)
    if m2:
        return int(m2.group(1))
    return 0


def infer_heading_from_text(text: str) -> int:
    """For flat-style docs (all Normal), infer heading level from content patterns.

    Returns pseudo-heading level (1-3) or 0 if not a heading.
    Heuristics:
      - Very short line (<=6 chars) with CJK only -> likely H1/H2 topic label
      - Pattern like "1.17餐费财税处理大全" -> numbered topic title (H2)
      - Short line ending in specific suffixes -> H3 sub-topic
    """
    text = text.strip()
    if not text:
        return 0

    # Numbered topic pattern: "1.17餐费财税处理大全" or "2.27 企税汇缴申报扣除类易错点"
    if re.match(r"^\d+\.\d+\s*[\u4e00-\u9fff]", text) and len(text) <= 30:
        return 2

    # Pure CJK short label (2-6 chars) -> likely a section heading
    if re.match(r"^[\u4e00-\u9fff]{2,8}$", text):
        return 1

    # Slightly longer CJK-only label (7-12 chars) without punctuation
    if re.match(r"^[\u4e00-\u9fff]{7,14}$", text) and len(text) <= 14:
        return 2

    return 0


# ── Text extraction ──────────────────────────────────────────────────────

def extract_docx_structured(file_path: Path) -> list[dict]:
    """Extract structured sections from a .docx file.

    Returns a list of section dicts:
      {"heading": str, "level": int, "content": [str], "subsections": [dict]}
    """
    import docx

    try:
        doc = docx.Document(str(file_path))
    except Exception as e:
        print(f"  WARN: Cannot open {file_path.name}: {e}")
        return []

    paragraphs = doc.paragraphs
    tables = doc.tables

    # Determine if this file uses heading styles
    has_heading_styles = any(
        get_heading_level(p.style.name) > 0 for p in paragraphs[:100]
    )

    sections = []
    current_section = None

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue

        style_level = get_heading_level(p.style.name)
        inferred_level = 0 if has_heading_styles else infer_heading_from_text(text)
        level = style_level or inferred_level

        if level > 0:
            # Start a new section
            section = {
                "heading": text,
                "level": level,
                "content": [],
            }
            sections.append(section)
            current_section = section
        else:
            # Append content to current section
            if current_section is not None:
                current_section["content"].append(text)
            else:
                # Content before any heading -> create implicit root section
                current_section = {
                    "heading": file_path.stem,
                    "level": 0,
                    "content": [text],
                }
                sections.append(current_section)

    # Extract table content and attach to last section
    for table in tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows and sections:
            sections[-1].setdefault("table_content", [])
            sections[-1]["table_content"].extend(table_rows)

    return sections


def extract_doc_structured(file_path: Path) -> list[dict]:
    """Extract structured sections from a .doc file via textutil conversion."""
    try:
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name

        result = subprocess.run(
            ["textutil", "-convert", "txt", "-output", tmp_path, str(file_path)],
            capture_output=True, timeout=30,
        )

        if result.returncode != 0:
            print(f"  WARN: textutil failed for {file_path.name}: rc={result.returncode}")
            return []

        text = Path(tmp_path).read_text(encoding="utf-8", errors="replace").strip()
        os.unlink(tmp_path)

        if not text:
            return []

        # Parse text line by line, using inferred heading detection
        lines = text.split("\n")
        sections = []
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            level = infer_heading_from_text(line)
            if level > 0:
                section = {"heading": line, "level": level, "content": []}
                sections.append(section)
                current_section = section
            else:
                if current_section is not None:
                    current_section["content"].append(line)
                else:
                    current_section = {
                        "heading": file_path.stem,
                        "level": 0,
                        "content": [line],
                    }
                    sections.append(current_section)

        return sections

    except subprocess.TimeoutExpired:
        print(f"  WARN: textutil timeout for {file_path.name}")
        return []
    except Exception as e:
        print(f"  WARN: extract_doc error for {file_path.name}: {e}")
        return []


# ── Graph node mapping ───────────────────────────────────────────────────

def make_node_id(category: str, heading: str) -> str:
    """Generate stable node ID from category + heading."""
    raw = f"{category}::{heading}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def has_risk_signal(text: str) -> bool:
    """Check if text contains risk-related keywords."""
    return any(kw in text for kw in RISK_KEYWORDS)


def classify_section(
    section: dict, category: str, base_node_type: str
) -> dict:
    """Classify a section into a graph node record.

    Returns a dict ready for JSON output with graph-aligned fields.
    """
    heading = section["heading"]
    content_lines = section.get("content", [])
    table_lines = section.get("table_content", [])
    full_content = "\n".join(content_lines)
    level = section["level"]

    # Determine node type: override to RiskIndicator if risk signals present
    combined_text = heading + " " + full_content
    if has_risk_signal(combined_text):
        node_type = "RiskIndicator"
    else:
        node_type = base_node_type

    node = {
        "id": make_node_id(category, heading),
        "node_type": node_type,
        "category": category,
        "name": _clean_heading(heading),
        "heading_level": level,
        "description": full_content[:3000],  # cap at 3000 chars
        "content_lines": len(content_lines),
        "source_file": "",  # filled by caller
    }

    # Add type-specific fields
    if node_type == "OP_BusinessScenario":
        node["trigger_condition"] = _extract_trigger(content_lines)
        node["policy_refs"] = _extract_policy_refs(content_lines)

    elif node_type == "OP_StandardCase":
        node["correct_treatment"] = _extract_treatment(content_lines)
        node["common_mistake"] = _extract_mistakes(content_lines)
        node["policy_refs"] = _extract_policy_refs(content_lines)

    elif node_type == "RiskIndicator":
        node["risk_points"] = _extract_risk_points(content_lines)
        node["prevention_measures"] = _extract_prevention(content_lines)

    elif node_type == "TaxOptimization":
        node["strategy"] = _extract_strategy(content_lines)
        node["policy_refs"] = _extract_policy_refs(content_lines)

    if table_lines:
        node["table_data"] = table_lines[:50]  # cap table rows

    return node


def _clean_heading(heading: str) -> str:
    """Remove leading number prefixes like '1.17' or '3.27'."""
    return re.sub(r"^\d+\.\d+\s*", "", heading).strip()


def _extract_policy_refs(lines: list[str]) -> list[str]:
    """Extract policy/law references from content lines."""
    refs = []
    patterns = [
        r"《.+?》",  # Chinese book title marks
        r"(?:国税|财税|税总|国发)\s*[\[【（(]\d{4}[\]】）)]\s*\d+\s*号",
    ]
    for line in lines:
        for pat in patterns:
            found = re.findall(pat, line)
            refs.extend(found)
    return list(dict.fromkeys(refs))[:20]  # dedupe, cap at 20


def _extract_trigger(lines: list[str]) -> str:
    """Extract trigger condition text."""
    trigger_kw = ["条件", "适用", "触发", "前提", "适用范围", "对象"]
    for i, line in enumerate(lines):
        if any(kw in line for kw in trigger_kw) and len(line) < 60:
            # Return the next few lines as context
            context = lines[i:i + 3]
            return "\n".join(context)
    return ""


def _extract_treatment(lines: list[str]) -> str:
    """Extract correct accounting treatment (journal entries, etc.)."""
    treatment_lines = []
    capture = False
    for line in lines:
        if any(kw in line for kw in ["借:", "贷:", "借：", "贷：", "会计分录", "账务处理"]):
            capture = True
        if capture:
            treatment_lines.append(line)
            if len(treatment_lines) >= 10:
                break
        if capture and not line.strip():
            break
    return "\n".join(treatment_lines)


def _extract_mistakes(lines: list[str]) -> str:
    """Extract common mistakes section."""
    mistake_kw = ["易错", "常见错误", "误区", "注意", "常见问题"]
    for i, line in enumerate(lines):
        if any(kw in line for kw in mistake_kw):
            context = lines[i:i + 5]
            return "\n".join(context)
    return ""


def _extract_risk_points(lines: list[str]) -> list[str]:
    """Extract risk point descriptions."""
    risk_kw = ["风险点", "风险", "隐患", "违规"]
    points = []
    for i, line in enumerate(lines):
        if any(kw in line for kw in risk_kw):
            points.append(line)
            # Add following context lines
            for j in range(1, 3):
                if i + j < len(lines) and lines[i + j].strip():
                    points.append(lines[i + j])
    return points[:20]


def _extract_prevention(lines: list[str]) -> list[str]:
    """Extract prevention/mitigation measures."""
    prev_kw = ["防范", "措施", "对策", "建议", "应对"]
    measures = []
    capture = False
    for line in lines:
        if any(kw in line for kw in prev_kw):
            capture = True
            continue
        if capture and line.strip():
            measures.append(line)
            if len(measures) >= 10:
                break
        elif capture and not line.strip():
            capture = False
    return measures


def _extract_strategy(lines: list[str]) -> str:
    """Extract tax optimization strategy description."""
    strategy_kw = ["方案", "策略", "优化", "筹划", "节税", "降低"]
    for i, line in enumerate(lines):
        if any(kw in line for kw in strategy_kw):
            context = lines[max(0, i - 1):i + 5]
            return "\n".join(context)
    # Fallback: return first few meaningful lines
    meaningful = [l for l in lines if len(l) > 10][:5]
    return "\n".join(meaningful)


# ── File processing ──────────────────────────────────────────────────────

def process_directory_file(
    file_path: Path, category: str, base_node_type: str
) -> list[dict]:
    """Process a single file and return a list of graph node dicts."""
    ext = file_path.suffix.lower()

    if ext == ".docx":
        sections = extract_docx_structured(file_path)
    elif ext == ".doc":
        sections = extract_doc_structured(file_path)
    else:
        return []

    if not sections:
        return []

    nodes = []
    rel_path = str(file_path.relative_to(file_path.parents[4]))  # relative to doc-tax/...

    # Filter: skip very short sections (likely noise) and level-0 root placeholders
    for section in sections:
        content_len = len("\n".join(section.get("content", [])))
        if content_len < 20 and section["level"] == 0:
            continue

        node = classify_section(section, category, base_node_type)
        node["source_file"] = rel_path
        nodes.append(node)

    return nodes


def process_all(base_dir: Path) -> dict[str, list[dict]]:
    """Process all priority files and directories.

    Returns a dict mapping output filename -> list of node dicts.
    """
    mindmap_dir = base_dir / MINDMAP_SUBDIR
    if not mindmap_dir.exists():
        print(f"ERROR: Mindmap directory not found: {mindmap_dir}")
        sys.exit(1)

    results = {}
    stats = {"files_processed": 0, "files_skipped": 0, "nodes_total": 0}

    # Process each category directory
    for dir_name, node_type in CATEGORY_NODE_MAP.items():
        cat_dir = mindmap_dir / dir_name
        if not cat_dir.exists():
            print(f"  WARN: Directory not found: {dir_name}")
            continue

        print(f"\nProcessing category: {dir_name} -> {node_type}")
        cat_nodes = []

        # Find all doc/docx files in this directory
        files = sorted(
            f for f in cat_dir.iterdir()
            if f.suffix.lower() in (".doc", ".docx") and not f.name.startswith("~")
        )

        for file_path in files:
            print(f"  {file_path.name}")
            nodes = process_directory_file(file_path, dir_name, node_type)
            if nodes:
                cat_nodes.extend(nodes)
                stats["files_processed"] += 1
            else:
                stats["files_skipped"] += 1

        if cat_nodes:
            # Sanitize directory name for output filename
            out_key = dir_name
            results[out_key] = cat_nodes
            stats["nodes_total"] += len(cat_nodes)
            print(f"  OK: {len(cat_nodes)} nodes extracted")

    # Process standalone files (directly under 分类合集/)
    for file_stem, node_type in STANDALONE_NODE_MAP.items():
        # Try both .docx and .doc
        for ext in (".docx", ".doc"):
            file_path = mindmap_dir / f"{file_stem}{ext}"
            if file_path.exists():
                print(f"\nProcessing standalone: {file_path.name} -> {node_type}")
                nodes = process_directory_file(file_path, file_stem, node_type)
                if nodes:
                    results[file_stem] = nodes
                    stats["files_processed"] += 1
                    stats["nodes_total"] += len(nodes)
                    print(f"  OK: {len(nodes)} nodes extracted")
                else:
                    stats["files_skipped"] += 1
                break

    return results, stats


# ── Main ─────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Extract structured content from 财税脑图3.0 mindmap docs"
    )
    parser.add_argument(
        "--base-dir",
        default="/Users/mauricewen/Projects/cognebula-enterprise",
        help="Project root directory",
    )
    args = parser.parse_args()

    base_dir = Path(args.base_dir).resolve()
    output_dir = base_dir / "data" / "extracted" / "mindmap"
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"NOTE: Base directory: {base_dir}")
    print(f"NOTE: Output directory: {output_dir}")
    print(f"NOTE: Timestamp: {datetime.now().isoformat()}")

    results, stats = process_all(base_dir)

    if not results:
        print("\nERROR: No content extracted")
        return 1

    # Write per-category JSON files
    for category, nodes in results.items():
        out_file = output_dir / f"{category}.json"
        out_file.write_text(
            json.dumps(nodes, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"  -> {out_file.name}: {len(nodes)} nodes")

    # Write combined output
    all_nodes = []
    for nodes in results.values():
        all_nodes.extend(nodes)

    combined_file = output_dir / "all_mindmap_nodes.json"
    combined_file.write_text(
        json.dumps(all_nodes, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    # Write extraction summary
    summary = {
        "extracted_at": datetime.now().isoformat(),
        "base_dir": str(base_dir),
        "files_processed": stats["files_processed"],
        "files_skipped": stats["files_skipped"],
        "total_nodes": stats["nodes_total"],
        "categories": {
            cat: {
                "node_count": len(nodes),
                "node_types": list(set(n["node_type"] for n in nodes)),
            }
            for cat, nodes in results.items()
        },
        "node_type_distribution": {},
    }

    # Count by node type
    type_counts = {}
    for node in all_nodes:
        nt = node["node_type"]
        type_counts[nt] = type_counts.get(nt, 0) + 1
    summary["node_type_distribution"] = type_counts

    summary_file = output_dir / "extraction_summary.json"
    summary_file.write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    # Print summary
    print(f"\n{'=' * 60}")
    print(f"OK: Extraction complete")
    print(f"  Files processed: {stats['files_processed']}")
    print(f"  Files skipped:   {stats['files_skipped']}")
    print(f"  Total nodes:     {stats['nodes_total']}")
    print(f"  Categories:      {len(results)}")
    print(f"\nNode type distribution:")
    for nt, count in sorted(type_counts.items(), key=lambda x: -x[1]):
        print(f"  {nt}: {count}")
    print(f"\nOutput directory: {output_dir}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
