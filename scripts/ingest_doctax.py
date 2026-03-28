#!/usr/bin/env python3
"""Parse doc-tax knowledge files and ingest into KuzuDB via API.

Uses Gemini 2.5 Pro for intelligent document structuring.
Processes: docx, pdf, xlsx -> structured JSON -> KG API ingest.

Usage:
    .venv/bin/python3 scripts/ingest_doctax.py [--dry-run] [--limit N]
"""
import hashlib
import json
import os
import sys
import time
import urllib.request
from pathlib import Path

# Config
PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_TAX_DIR = PROJECT_ROOT / "doc-tax"
OUTPUT_DIR = PROJECT_ROOT / "data" / "extracted"
KG_API = os.environ.get("KG_API", "http://100.75.77.112:8400")
GEMINI_MODEL = "gemini-3.1-pro"  # via Poe API
BATCH_SIZE = 30
DRY_RUN = "--dry-run" in sys.argv
LIMIT = None
for i, arg in enumerate(sys.argv):
    if arg == "--limit" and i + 1 < len(sys.argv):
        LIMIT = int(sys.argv[i + 1])

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def get_api_key():
    key = os.environ.get("GEMINI_API_KEY")
    if key:
        return key
    for p in [Path.home() / ".openclaw" / ".env", Path(".env")]:
        if p.exists():
            for line in p.read_text().splitlines():
                if line.startswith("GEMINI_API_KEY="):
                    return line.split("=", 1)[1].strip()
    return None


API_KEY = get_api_key()
if not API_KEY:
    print("ERROR: GEMINI_API_KEY not found")
    sys.exit(1)


def make_id(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()[:16]


def extract_docx(path: Path) -> str:
    """Extract text from a docx file."""
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[EXTRACT ERROR: {e}]"


def extract_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for page in reader.pages[:30]:  # Max 30 pages
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return "\n\n".join(pages)
    except Exception as e:
        return f"[EXTRACT ERROR: {e}]"


def extract_xlsx(path: Path) -> str:
    """Extract text from Excel file (first sheet, first 200 rows)."""
    try:
        import csv
        import subprocess
        # Use xlsx2csv if available, otherwise basic extraction
        result = subprocess.run(
            ["python3", "-c", f"""
import openpyxl
wb = openpyxl.load_workbook('{path}', read_only=True, data_only=True)
ws = wb.active
rows = []
for i, row in enumerate(ws.iter_rows(values_only=True)):
    if i > 200: break
    cells = [str(c or '') for c in row]
    rows.append(' | '.join(cells))
print('\\n'.join(rows))
"""],
            capture_output=True, text=True, timeout=30
        )
        return result.stdout[:5000] if result.returncode == 0 else f"[XLSX ERROR: {result.stderr[:200]}]"
    except Exception as e:
        return f"[EXTRACT ERROR: {e}]"


CHUNK_SIZE = 3000  # smaller chunks = more focused extraction

STRUCTURE_PROMPT = """你是中国财税知识图谱的知识提取引擎。你的任务是从文档中提取每一条可独立查询的知识原子。

## 文档: {filename} | 分类: {category} | 片段 {chunk_idx}/{total_chunks}

{text}

## 提取为 JSON 数组。每个元素必须包含：
- title: 精确标题，必须包含核心数字或关键词（如"小规模纳税人增值税征收率3%"而非"增值税征收率"）
- content: 完整知识描述，保留所有数字、条件、公式、例外情况

## 必须提取的知识原子类型（每看到一个就提取一个节点）：
1. **税率**: "XX税 Y% 适用于 Z 场景" — 每个税率一个节点
2. **起征点/阈值**: "年销售额超过500万 → 一般纳税人" — 每个阈值一个节点
3. **计算公式**: "应纳税额 = 销项税额 - 进项税额" — 每个公式一个节点
4. **申报填写**: "XX表第N行填写YY数据" — 每个填写规则一个节点
5. **优惠政策**: "XX情况下减免/免征/减半" — 每条优惠一个节点
6. **处罚/风险**: "未按时申报 → 滞纳金/罚款" — 每条处罚一个节点
7. **时间节点**: "XX税每月/季度N日前申报" — 每个截止日一个节点
8. **会计分录**: "借:XX 贷:YY" — 每个分录一个节点
9. **操作步骤**: 流程中每个步骤单独一个节点
10. **对比区别**: "A vs B 的区别在于..." — 每组对比一个节点
11. **具体案例**: "某企业XX情况下如何处理" — 每个案例一个节点
12. **注意事项**: "特别注意：XX" — 每条注意事项一个节点

## 反例（不要这样做）：
BAD: {{"title": "增值税概述", "content": "增值税是一种流转税..."}} — 太泛
GOOD: {{"title": "一般纳税人增值税基本税率13%适用货物销售", "content": "一般纳税人销售货物（非特殊商品）适用13%增值税税率。计算方式：销项税额=不含税销售额×13%。发票类型：增值税专用发票或普通发票。"}}

目标：从这段内容提取 15-40 个知识原子。宁多勿少，宁细勿粗。"""


def _timeout_handler(signum, frame):
    raise TimeoutError("Gemini API call timed out")


sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from llm_client import llm_generate_json


def gemini_call(prompt: str) -> list[dict]:
    """Call LLM via Poe API, returns parsed JSON list."""
    import signal
    old_handler = signal.signal(signal.SIGALRM, _timeout_handler)
    signal.alarm(90)
    try:
        result = llm_generate_json(prompt, model=GEMINI_MODEL, temperature=0.1, max_tokens=16384)
        if isinstance(result, list):
            return result
        return []
    except TimeoutError:
        print("    LLM TIMEOUT (90s)")
        return []
    except Exception as e:
        print(f"    LLM error: {e}")
        return []
    finally:
        signal.alarm(0)
        signal.signal(signal.SIGALRM, old_handler)


def split_chunks(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    """Split text into chunks at paragraph boundaries."""
    paragraphs = text.split("\n")
    chunks = []
    current = []
    current_len = 0
    for para in paragraphs:
        if current_len + len(para) > chunk_size and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(para)
        current_len += len(para) + 1
    if current:
        chunks.append("\n".join(current))
    return [c for c in chunks if len(c.strip()) > 50]


def gemini_structure(text: str, filename: str, category: str) -> list[dict]:
    """Use Gemini 2.5 Pro to structure document text into KG nodes (chunked)."""
    if not text or len(text) < 30 or text.startswith("[EXTRACT ERROR"):
        return []

    chunks = split_chunks(text, CHUNK_SIZE)
    if not chunks:
        return []

    # Process all chunks — no cap (Gemini 2.5 Pro handles volume)
    all_nodes = []
    total = len(chunks)

    for i, chunk in enumerate(chunks):
        prompt = STRUCTURE_PROMPT.format(
            filename=filename,
            category=category,
            chunk_idx=i + 1,
            total_chunks=total,
            text=chunk,
        )
        nodes = gemini_call(prompt)
        all_nodes.extend(nodes)
        # Progress for large files
        if total > 5 and (i + 1) % 5 == 0:
            print(f"      chunk {i+1}/{total}: {len(all_nodes)} nodes so far")
        time.sleep(0.8)  # Gemini RPM management

    return all_nodes


def ingest_batch(table: str, nodes: list[dict]) -> tuple[int, int]:
    """Batch ingest via KG API."""
    if DRY_RUN:
        return len(nodes), 0

    inserted = 0
    errors = 0
    for i in range(0, len(nodes), BATCH_SIZE):
        batch = nodes[i:i + BATCH_SIZE]
        payload = json.dumps({"table": table, "nodes": batch}).encode()
        req = urllib.request.Request(
            f"{KG_API}/api/v1/ingest",
            data=payload,
            headers={"Content-Type": "application/json"},
        )
        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                r = json.loads(resp.read())
                inserted += r.get("inserted", 0)
                errors += r.get("errors", 0)
        except Exception as e:
            errors += len(batch)
    return inserted, errors


def get_category_from_path(path: Path) -> str:
    """Derive category from file path."""
    parts = path.relative_to(DOC_TAX_DIR).parts
    if len(parts) >= 2:
        return parts[0] + "/" + parts[1]
    return parts[0] if parts else "uncategorized"


def main():
    print(f"=== CogNebula doc-tax Ingest Pipeline ===")
    print(f"Source: {DOC_TAX_DIR}")
    print(f"API: {KG_API}")
    print(f"Model: {GEMINI_MODEL}")
    print(f"Dry run: {DRY_RUN}")
    print()

    # Load processed files tracker (skip already-done files)
    tracker_file = OUTPUT_DIR / "doctax_processed.json"
    processed = set()
    if tracker_file.exists():
        try:
            processed = set(json.loads(tracker_file.read_text()))
        except:
            pass

    # Collect all processable files
    files = []
    for ext in ["*.docx", "*.pdf"]:
        files.extend(DOC_TAX_DIR.rglob(ext))
    # Filter out temp files and already-processed
    files = [f for f in files if not f.name.startswith("~") and not f.name.startswith(".")
             and str(f) not in processed]
    files.sort(key=lambda f: f.stat().st_size)  # Small files first

    if LIMIT:
        files = files[:LIMIT]

    print(f"Files to process: {len(files)} (already done: {len(processed)})", flush=True)
    print(flush=True)

    # Track progress
    total_nodes = 0
    total_inserted = 0
    total_errors = 0
    all_extracted = []

    for i, fpath in enumerate(files):
        ext = fpath.suffix.lower()
        category = get_category_from_path(fpath)
        size_kb = fpath.stat().st_size / 1024

        # Skip very large files (> 20MB)
        if size_kb > 20000:
            print(f"  [{i+1}/{len(files)}] SKIP (too large: {size_kb:.0f}KB): {fpath.name}")
            continue

        print(f"  [{i+1}/{len(files)}] {ext} {size_kb:.0f}KB {category}/{fpath.name}")

        # Extract text
        if ext == ".docx":
            text = extract_docx(fpath)
        elif ext == ".pdf":
            text = extract_pdf(fpath)
        else:
            continue

        if len(text) < 30:
            print(f"    -> Skip (too short: {len(text)} chars)")
            continue

        # Structure with Gemini
        nodes = gemini_structure(text, fpath.name, category)
        if not nodes:
            print(f"    -> No nodes extracted", flush=True)
            # Mark as processed even if no nodes (avoid re-trying timeout files)
            processed.add(str(fpath))
            tracker_file.write_text(json.dumps(sorted(processed), ensure_ascii=False))
            continue

        # Prepare for ingest — write as LawOrRegulation (unified node type)
        ingest_nodes = []
        for node in nodes:
            nid = make_id(fpath.name + ":" + node.get("title", ""))
            ingest_nodes.append({
                "id": nid,
                "title": (node.get("title") or "")[:500],
                "fullText": (node.get("content") or "")[:5000],
                "sourceUrl": f"doc-tax/{fpath.name}",
                "regulationNumber": "",
                "regulationType": "doctax_knowledge",
                "source": (node.get("category") or category)[:200],
                "type": node.get("node_type", "knowledge"),
            })
            all_extracted.append({**node, "id": nid, "source_file": str(fpath.name)})

        # Ingest as LawOrRegulation (VPS ingest endpoint supports this table)
        ins, err = ingest_batch("LawOrRegulation", ingest_nodes)
        total_inserted += ins
        total_errors += err
        total_nodes += len(ingest_nodes)

        print(f"    -> {len(nodes)} nodes -> LawOrRegulation (ins:{ins} err:{err})", flush=True)

        # Track processed file
        processed.add(str(fpath))
        tracker_file.write_text(json.dumps(sorted(processed), ensure_ascii=False))

        # Rate limit (Gemini RPM)
        time.sleep(1.5)

    # Save all extracted data (append to existing)
    output_file = OUTPUT_DIR / "doctax_extracted.json"
    existing = []
    if output_file.exists():
        try:
            with open(output_file) as f:
                existing = json.load(f)
        except:
            pass
    # Dedup by id
    seen_ids = {n.get("id") for n in existing}
    for n in all_extracted:
        if n.get("id") not in seen_ids:
            existing.append(n)
            seen_ids.add(n.get("id"))
    with open(output_file, "w") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)

    print()
    print(f"=== DONE ===")
    print(f"Files processed: {len(files)}")
    print(f"Nodes extracted: {total_nodes}")
    print(f"Inserted: {total_inserted}")
    print(f"Errors: {total_errors}")
    print(f"Saved to: {output_file}")


if __name__ == "__main__":
    main()
