#!/usr/bin/env python3
"""Extract text from local doc-tax files (doc/docx/xlsx/pdf) into JSON for knowledge graph ingestion.

Walks the doc-tax directory tree, extracts text content, classifies by directory path,
and outputs a JSON array compatible with finance_tax_processor.py.

Supported formats:
- .docx: python-docx (native Python)
- .doc: macOS textutil (converts to txt via subprocess)
- .xlsx/.xls: openpyxl (sheet names + first N rows as text)
- .pdf: skipped (separate pipeline)

Usage:
    python extract_local_docs.py --input doc-tax --output data/raw/local-doctax
    python extract_local_docs.py --input doc-tax --output data/raw/local-doctax --xlsx-deep
"""

import argparse
import hashlib
import json
import os
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

# ── Category extraction from directory path ─────────────────────────────

# Map top-level directory prefix to category prefix
TOP_LEVEL_PREFIX = {
    "1.财税脑图分类合集": "脑图",
    "2.会计小白经验提升": "教程",
    "3.财税实战工具大全": "工具",
}


def extract_category(file_path: Path, base_dir: Path) -> str:
    """Derive category from the file's position in the directory tree.

    Examples:
        财税脑图3.0/1.财税脑图分类合集/税务稽查/foo.doc -> "脑图/税务稽查"
        财税脑图3.0/2.会计小白经验提升/2.纳税申报/foo.docx -> "教程/纳税申报"
        财税脑图3.0/3.财税实战工具大全/4.行业财务制度范本/foo.xlsx -> "工具/行业财务制度"
        企业报告示例/foo.pdf -> "企业报告"
    """
    try:
        rel = file_path.relative_to(base_dir)
    except ValueError:
        return "其他"

    parts = rel.parts  # e.g. ("财税脑图3.0", "1.财税脑图分类合集", "税务稽查", "file.doc")

    if len(parts) >= 1 and parts[0] == "企业报告示例":
        return "企业报告"

    # Files directly under 财税脑图3.0/
    if len(parts) >= 2 and parts[0].startswith("财税脑图"):
        top_dir = parts[1]  # e.g. "1.财税脑图分类合集"
        prefix = TOP_LEVEL_PREFIX.get(top_dir, "")

        if not prefix:
            # File directly under 财税脑图3.0/ (e.g. 财税脑图.docx)
            return "脑图"

        if len(parts) >= 4:
            # File inside a sub-category directory
            sub = parts[2]
            import re
            sub_clean = re.sub(r"^\d+[.．、]", "", sub)
            return f"{prefix}/{sub_clean}"

        if len(parts) == 3:
            # File directly in top-level category dir (e.g. 1.财税脑图分类合集/税费相关.docx)
            # The "sub" is the filename itself; derive category from stem
            stem = file_path.stem
            import re
            stem_clean = re.sub(r"^\d+[.．、]\s*", "", stem)
            return f"{prefix}/{stem_clean}"

        # File directly in top-level category dir with no sub
        return prefix

    return "其他"


# ── Text extraction ─────────────────────────────────────────────────────

def extract_docx(file_path: Path) -> str:
    """Extract text from .docx using python-docx."""
    try:
        import docx
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[EXTRACT_ERROR: {e}]"


def extract_doc(file_path: Path) -> str:
    """Extract text from .doc using macOS textutil (converts to txt)."""
    try:
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-output", tmp_path, str(file_path)],
            capture_output=True, timeout=30,
        )
        if result.returncode == 0 and os.path.exists(tmp_path):
            text = Path(tmp_path).read_text(encoding="utf-8", errors="replace")
            os.unlink(tmp_path)
            return text.strip()
        os.unlink(tmp_path) if os.path.exists(tmp_path) else None
        return f"[TEXTUTIL_FAILED: rc={result.returncode}]"
    except subprocess.TimeoutExpired:
        return "[TEXTUTIL_TIMEOUT]"
    except Exception as e:
        return f"[EXTRACT_ERROR: {e}]"


def extract_xlsx(file_path: Path, max_rows: int = 20) -> str:
    """Extract sheet names and first N rows from .xlsx/.xls."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"[Sheet: {sheet_name}]")
            row_count = 0
            for row in ws.iter_rows(max_row=max_rows, values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line = " | ".join(cells)
                if line.strip() and line.strip(" |"):
                    parts.append(line)
                    row_count += 1
            if row_count == 0:
                parts.append("(empty sheet)")
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        return f"[XLSX_ERROR: {e}]"


def extract_xlsx_deep(file_path: Path) -> dict | None:
    """Deep extraction for specific high-value xlsx files.

    Returns structured data dict or None if not a target file.
    """
    name = file_path.name

    if "税负率" in name and "计算" in name:
        return _extract_tax_burden_rates(file_path)
    elif "预警指标" in name:
        return _extract_warning_indicators(file_path)
    elif "毛利率" in name and "税负" in name:
        return _extract_gross_margin_tax(file_path)

    return None


def _extract_tax_burden_rates(file_path: Path) -> dict:
    """Extract industry tax burden rates from 各行业税负率及计算表.xlsx."""
    import openpyxl
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    records = []
    for ws in wb.worksheets:
        header = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            if i == 0 or any(k in "".join(cells) for k in ("行业", "税负", "Industry")):
                header = cells
                continue
            if header and any(c.strip() for c in cells):
                record = {}
                for j, val in enumerate(cells):
                    if j < len(header) and header[j].strip():
                        record[header[j].strip()] = val.strip()
                if record:
                    records.append(record)
    wb.close()
    return {
        "file": file_path.name,
        "type": "industry_tax_burden_rates",
        "records": records[:200],  # cap at 200 rows
    }


def _extract_warning_indicators(file_path: Path) -> dict:
    """Extract tax warning indicators from 税务预警指标测算系统.xlsx."""
    import openpyxl
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    indicators = []
    for ws in wb.worksheets:
        header = None
        for i, row in enumerate(ws.iter_rows(max_row=100, values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            joined = "".join(cells)
            if any(k in joined for k in ("指标", "预警", "阈值", "公式")):
                if header is None:
                    header = cells
                    continue
            if header and any(c.strip() for c in cells):
                record = {}
                for j, val in enumerate(cells):
                    if j < len(header) and header[j].strip():
                        record[header[j].strip()] = val.strip()
                if record:
                    indicators.append(record)
    wb.close()
    return {
        "file": file_path.name,
        "type": "tax_warning_indicators",
        "indicators": indicators[:100],
    }


def _extract_gross_margin_tax(file_path: Path) -> dict:
    """Extract gross margin to tax burden relationships."""
    import openpyxl
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    relationships = []
    for ws in wb.worksheets:
        header = None
        for i, row in enumerate(ws.iter_rows(max_row=100, values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            joined = "".join(cells)
            if any(k in joined for k in ("毛利", "税负", "税率")):
                if header is None:
                    header = cells
                    continue
            if header and any(c.strip() for c in cells):
                record = {}
                for j, val in enumerate(cells):
                    if j < len(header) and header[j].strip():
                        record[header[j].strip()] = val.strip()
                if record:
                    relationships.append(record)
    wb.close()
    return {
        "file": file_path.name,
        "type": "gross_margin_tax_relationship",
        "relationships": relationships[:100],
    }


# ── Main processing ─────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".doc", ".docx", ".xlsx", ".xls"}


def make_doc_id(file_path: Path) -> str:
    """Generate stable ID from file path (sha256)."""
    return hashlib.sha256(str(file_path).encode("utf-8")).hexdigest()[:16]


def process_file(file_path: Path, base_dir: Path) -> dict | None:
    """Process a single file and return a document dict."""
    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return None

    title = file_path.stem
    # Strip leading version numbers like "2.27 " or "3.83 "
    import re
    title_clean = re.sub(r"^\d+\.\d+\s+", "", title)

    category = extract_category(file_path, base_dir)

    # Extract text based on file type
    if ext == ".docx":
        content = extract_docx(file_path)
    elif ext == ".doc":
        content = extract_doc(file_path)
    elif ext in (".xlsx", ".xls"):
        content = extract_xlsx(file_path)
    else:
        return None

    if not content or content.startswith("[EXTRACT_ERROR"):
        return None

    # Truncate content to 5000 chars
    content_truncated = content[:5000]

    return {
        "id": make_doc_id(file_path),
        "title": title_clean,
        "content": content_truncated,
        "source": "local_doctax",
        "type": category,
        "url": "",
        "date": "",
        "crawled_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
    }


def main():
    parser = argparse.ArgumentParser(description="Extract text from local doc-tax files")
    parser.add_argument("--input", required=True, help="Input doc-tax directory")
    parser.add_argument("--output", required=True, help="Output directory for JSON")
    parser.add_argument("--xlsx-deep", action="store_true",
                        help="Also extract structured data from key xlsx files")
    args = parser.parse_args()

    input_dir = Path(args.input).resolve()
    output_dir = Path(args.output).resolve()

    if not input_dir.exists():
        print(f"ERROR: Input directory not found: {input_dir}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    # Collect all supported files
    all_files = sorted(f for f in input_dir.rglob("*") if f.suffix.lower() in SUPPORTED_EXTENSIONS)
    print(f"NOTE: Found {len(all_files)} supported files in {input_dir}")

    # Count by extension
    ext_counts = {}
    for f in all_files:
        ext = f.suffix.lower()
        ext_counts[ext] = ext_counts.get(ext, 0) + 1
    for ext, count in sorted(ext_counts.items()):
        print(f"  {ext}: {count}")

    # Process all files
    docs = []
    errors = []
    deep_data = []

    for i, file_path in enumerate(all_files, 1):
        if i % 100 == 0:
            print(f"  Processing {i}/{len(all_files)}...")

        try:
            doc = process_file(file_path, input_dir)
            if doc:
                docs.append(doc)
            else:
                errors.append(f"SKIP: {file_path.name}")
        except Exception as e:
            errors.append(f"ERROR: {file_path.name}: {e}")

        # Deep xlsx extraction
        if args.xlsx_deep and file_path.suffix.lower() in (".xlsx", ".xls"):
            try:
                deep = extract_xlsx_deep(file_path)
                if deep:
                    deep_data.append(deep)
            except Exception as e:
                errors.append(f"DEEP_ERROR: {file_path.name}: {e}")

    # Write main output
    out_file = output_dir / "local_docs.json"
    out_file.write_text(json.dumps(docs, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nOK: Extracted {len(docs)} documents -> {out_file}")

    # Write deep xlsx data if any
    if deep_data:
        deep_file = output_dir / "xlsx_structured.json"
        deep_file.write_text(json.dumps(deep_data, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"OK: Extracted {len(deep_data)} structured xlsx datasets -> {deep_file}")

    # Write errors log
    if errors:
        err_file = output_dir / "extract_errors.log"
        err_file.write_text("\n".join(errors), encoding="utf-8")
        print(f"WARN: {len(errors)} errors/skips logged -> {err_file}")

    # Category distribution
    cat_counts = {}
    for d in docs:
        cat_counts[d["type"]] = cat_counts.get(d["type"], 0) + 1
    print("\nCategory distribution:")
    for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1]):
        print(f"  {cat}: {count}")

    return 0 if docs else 1


if __name__ == "__main__":
    sys.exit(main())
